import streamlit as st    #Creates the web app interface.
from docx import Document  
from docx.shared import Inches, Pt  #Generates and formats the Word (.docx) file.
import fitz  # PyMuPDF  #Extracts text and images from PDF pages
import tempfile
import os
import io
import time #Used for progress bar animation and delays.

# --------------------------- #
# PAGE CONFIGURATION
# --------------------------- #

st.set_page_config(page_title="PDF → Word Converter", page_icon="📄")

# --------------------------- #
# APP TITLE
# --------------------------- #

st.title("📄 PDF → Word Converter with Images")
st.write("Convert your PDF into an editable Word document (including text and line formatting).")

# --------------------------- #
# FILE UPLOADER
# --------------------------- #

uploaded_file = st.file_uploader("Upload a PDF file", type=["pdf"])

if uploaded_file:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
        temp_pdf.write(uploaded_file.read())
        pdf_path = temp_pdf.name

    if st.button("🚀 Convert to Word"):
        progress = st.progress(0)
        st.info("Converting your PDF to Word (text + images)...")
        
        try:
            for i in range(0, 100, 20):
                time.sleep(0.1)
                progress.progress(i + 10)

            # Open PDF that one uploaded
            pdf_doc = fitz.open(pdf_path)
            word_doc = Document()

            for page_num, page in enumerate(pdf_doc, start=1):
                text = page.get_text("text") ########### Read text from PDF ##########

                # Split text into lines to preserve spacing
                lines = text.split('\n')
                for line in lines:
                    clean_line = ''.join(ch for ch in line if ch.isprintable()).replace('\x00', '').strip()
                    if clean_line:
                        word_doc.add_paragraph(clean_line)
                word_doc.add_paragraph("")  # extra line break between pages

                # Extract and add images
                image_list = page.get_images(full=True)
                if image_list:
                    word_doc.add_paragraph("📸 Images on this page:")
                    for img_index, img in enumerate(image_list, start=1):
                        xref = img[0]
                        base_image = pdf_doc.extract_image(xref)
                        image_bytes = base_image["image"]

                        image_stream = io.BytesIO(image_bytes)
                        try:
                            word_doc.add_picture(image_stream, width=Inches(5.5)) ############ Page separation ############
                        except Exception as e:
                            st.warning(f"⚠️ Could not add image {img_index} from page {page_num}: {e}")

                word_doc.add_page_break()

            # Save and download final Word file
            output_path = pdf_path.replace(".pdf", ".docx")
            word_doc.save(output_path)

            with open(output_path, "rb") as f:
                st.success("✅ Conversion complete (Text format preserved)!")
                st.download_button(
                    label="📥 Download Word File",
                    data=f,
                    file_name="converted_with_formatting.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        except Exception as e:
            st.error(f"❌ Error: {e}")

        finally:
            try:
                if os.path.exists(pdf_path):
                    os.remove(pdf_path)
                if os.path.exists(output_path):
                    os.remove(output_path)
            except:
                pass
